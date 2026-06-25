# -*- coding: utf-8 -*-
"""Genera el resumen de avances (DOCX) para enviar a dirección."""
import tempfile, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x0C, 0x1E, 0x35)
BLUE = RGBColor(0x1A, 0x8E, 0xE6)
GREEN = RGBColor(0x22, 0xC4, 0x68)
GRAY = RGBColor(0x4B, 0x64, 0x80)

doc = Document()

# Estilo base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x22, 0x2A, 0x33)

def title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = NAVY
    p.space_after = Pt(2)
    return p

def subtitle(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11); r.font.color.rgb = GRAY
    p.space_after = Pt(10)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.space_before = Pt(12); p.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(13.5); r.font.bold = True; r.font.color.rgb = BLUE
    return p

def bullet(parts):
    """parts: str o lista de (texto, bold)"""
    p = doc.add_paragraph(style="List Bullet")
    p.space_after = Pt(3)
    if isinstance(parts, str):
        parts = [(parts, False)]
    for txt, bold in parts:
        r = p.add_run(txt)
        r.font.size = Pt(11); r.font.bold = bold
        r.font.color.rgb = NAVY if bold else RGBColor(0x33, 0x3A, 0x42)
    return p

def hr():
    p = doc.add_paragraph()
    r = p.add_run("─" * 60)
    r.font.color.rgb = RGBColor(0xC8, 0xDE, 0xF0)
    p.space_after = Pt(6)

# ── Encabezado ──
title("Resumen de avances — Sitio web LocalExpertiz")
subtitle("Fecha: 22 de junio, 2026")
hr()

# 1
h2("1. Reposicionamiento estratégico: de redes sociales → diseño web")
doc.add_paragraph(
    "Se reorientó todo el sitio para que comunique el nuevo negocio: la construcción de "
    "páginas web (antes estaba enfocado en redes sociales)."
)
bullet("Actualizados: título y SEO, hero, banda de servicios, sección “Nosotros” y todos los textos clave.")
bullet([("Se eliminaron las 6 páginas internas de redes sociales", True), (" y se corrigieron todos los enlaces (sin enlaces rotos).", False)])
bullet([("El blog", False), (" (que era 100% de redes sociales) quedó en estado ", False), ("“Próximamente”", True), (" con enfoque de diseño web.", False)])

# 2
h2("2. Investigación de mercado (entregable: presentación)")
doc.add_paragraph(
    "Se realizó una investigación de competidores enfocada en construcción de páginas web, "
    "con precios reales de México y EE.UU., quién ofrece qué y a cuánto, y los modelos de cobro. "
    "Incluye una propuesta de 3 paquetes con precios definidos:"
)
bullet([("Web Lanzamiento", True), (" (pago único) — desde $18,000 MXN / $2,500 USD", False)])
bullet([("Web + Crecimiento", True), (" (con soporte mensual) — paquete recomendado", False)])
bullet([("Web Autónoma", True), (" (CMS editable por el cliente)", False)])
bullet([("Entregable: presentación ", False), ("LocalExpertiz-Investigacion-Web.pptx", True), (" (18 diapositivas, más de 25 fuentes).", False)])

# 3
h2("3. Nuevos elementos en el sitio")
bullet([("Sección de Precios", True), (" nueva, con los 3 paquetes y un botón para cambiar entre MXN y USD.", False)])
bullet([("Servicios", True), (" reescritos a los 3 pilares de web: Diseño & Desarrollo, Soporte & Mantenimiento, y Webs Autoadministrables.", False)])
bullet([("Portafolio", True), (" actualizado a 4 proyectos de diseño web, cada uno con su página de detalle propia (Tienda en Línea, Landing Pages, Web Autoadministrable y Sitio Corporativo).", False)])

# 4
h2("4. Rediseño visual y experiencia (UX)")
bullet([("Tipografía mucho más grande", True), (" y moderna en todo el sitio.", False)])
bullet("Animaciones de aparición modernas, scroll suave, parallax en imágenes y barra de progreso de lectura.")
bullet("Encabezado transparente sobre el hero (se vuelve sólido al bajar).")
bullet("Sección de testimonios rediseñada (tarjetas apiladas con estrellas y rotación automática).")
bullet([("Palabra animada en el hero que alterna ", False), ("“MARCA” ↔ “EMPRESA”", True), (" automáticamente.", False)])
bullet([("Responsive", True), (" (móvil/tablet) revisado y afinado.", False)])

# 5
h2("5. Pendientes / decisiones para validar")
bullet([("Validar los precios", True), (" de los 3 paquetes contra nuestros costos y márgenes reales.", False)])
bullet([("Reemplazar imágenes y métricas del portafolio", True), (" (hoy son placeholder) por proyectos y resultados reales.", False)])
bullet([("Sumar testimonios reales", True), (" de clientes.", False)])
bullet([("Definir si reforzamos el verde de marca", True), (" (hoy domina el azul).", False)])

hr()
foot = doc.add_paragraph()
fr = foot.add_run("LocalExpertiz · Diseño y Desarrollo Web · USA + MX")
fr.font.size = Pt(9); fr.font.color.rgb = GRAY
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

out = os.path.join(tempfile.gettempdir(), "LocalExpertiz-Resumen-Avances.docx")
doc.save(out)
print("SAVED:", out)
print("SIZE_KB:", round(os.path.getsize(out) / 1024, 1))
